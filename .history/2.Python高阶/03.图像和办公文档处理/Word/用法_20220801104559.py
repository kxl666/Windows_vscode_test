import docx

doc = docx.Document('./File/test.docx')
# 1.读取文档
# 返回文档的paragraphs对象个数
print(len(doc.paragraphs))
# 获取文档的段落对象文本
print(doc.paragraphs[7].text)
# 段落的样式
# print(doc.paragraphs[7].style.name)
# style.font.name
# style.font.size
# style.font.bold
# style.font.italic
# style.font.underline
# 获取文档的段落对象样式

# runs是段落的字符串，每个字符串是一个run对象,一个段落可能存在多个字符串
print(len(doc.paragraphs[15].runs))
for i in range(len(doc.paragraphs[15].runs)):
    print(doc.paragraphs[15].runs[i].text)

# 2.从.docx 文件中取得完整的文本
fullText = [paragraph.text for paragraph in doc.paragraphs]
print('\n'.join(fullText))

# 3.Run属性
# 获取run的属性
print(doc.paragraphs[15].runs[0].text)  # 获取run的文本
print(doc.paragraphs[15].style.name)  # 获取run的样式
doc.paragraphs[15].style = 'Heading 1'  # 设置run的样式
doc.paragraphs[15].runs[0].text = '新的标题'  # 设置run的文本
print(doc.paragraphs[15].runs[0].text)
doc.paragraphs[15].runs[0].style.name = 'QuoteChar'  # 设置run的样式
doc.paragraphs[15].runs[0].underline = True  # 设置run的下划线
doc.paragraphs[15].runs[0].bold = True  # 设置run的粗体

# 4.写入文档
# 添加文本到段落
doc.add_paragraph('新的段落')
# 添加文本到段落，并设置样式
doc.add_paragraph('新的段落', style='Heading 1')
# 添加文本到段落，并设置样式，并设置下划线
doc.add_paragraph('新的段落', style='Heading 1').add_run('新的段落').underline = True
# 添加文本到段落，并设置样式，并设置粗体
doc.add_paragraph('新的段落', style='Heading 1').add_run('新的段落').bold = True

# # 添加标题0-4,表示不同标题层次
# doc.add_heading('新的标题', 0)
# # 添加标题，并设置样式
# doc.add_heading('新的标题', 1).style = 'Heading 1'
# # 添加标题，并设置样式，并设置下划线
# doc.add_heading('新的标题', 1).style = 'Heading 1'.add_run('新的段落').underline = True
# # 添加标题，并设置样式，并设置粗体
# doc.add_heading('新的标题', 1).style = 'Heading 1'.add_run('新的段落').bold = True

doc.save('./File/test02.docx')
