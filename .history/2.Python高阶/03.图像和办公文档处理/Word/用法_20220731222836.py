import docx

doc = docx.Document('./File/test.docx')
# 1.读取文档
# 返回文档的paragraphs对象个数
print(len(doc.paragraphs))
# 获取文档的段落对象文本
print(doc.paragraphs[7].text)
# 段落的样式
# print(doc.paragraphs[7].style.name)
# style.font.name
# style.font.size
# style.font.bold
# style.font.italic
# style.font.underline
# 获取文档的段落对象样式

# runs是段落的字符串，每个字符串是一个run对象
print(len(doc.paragraphs[15].runs))
for i in range(len(doc.paragraphs[15].runs)):
    print(doc.paragraphs[15].runs[i].text)

# 2.设置文档样式
# 设置字体样式
doc.paragraphs[15].runs[0].underline = True

doc.
