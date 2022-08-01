import docx

doc = docx.Document('./File/test.docx')
# 1.读取文档
# 返回文档的paragraphs对象个数
print(len(doc.paragraphs))
# 返回文档的tables对象个数
print(len(doc.tables))
# 返回文档的sections对象个数
print(len(doc.sections))
