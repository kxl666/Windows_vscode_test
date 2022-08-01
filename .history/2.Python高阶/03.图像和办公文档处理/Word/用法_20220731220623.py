import docx

doc = docx.Document('./File/test.docx')
# 1.读取文档
# 返回文档的paragraphs对象个数
print(len(doc.paragraphs))
# 获取文档的段落对象文本
print(doc.paragraphs[7].text)
# 段落的样式
print(doc.paragraphs[7].style.name)
# 获取文档的段落对象样式
print(doc.paragraphs[7].style.font.name)
# 获取文档的段落对象字体大小
print(doc.paragraphs[7].style.font.size)
# 获取文档的段落对象字体加粗
print(doc.paragraphs[7].style.font.bold)
# 获取文档的段落对象字体倾斜
print(doc.paragraphs[7].style.font.italic)
# 获取文档的段落对象字体下划线
print(doc.paragraphs[7].style.font.underline)

print(doc.)